import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Establece el color de fondo de una celda de tabla en Hexadecimal."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Establece márgenes internos para una celda."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_left_border(cell, color_hex="D4AF37", size="36"):
    """Agrega un borde izquierdo grueso a una celda para cajas de llamada / citas."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size)
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color_hex)
    tcBorders.append(left)
    
    for b_name in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{b_name}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
        
    tcPr.append(tcBorders)

def make_callout(doc, text, title=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F4F6F9")
    set_cell_left_border(cell, "D4AF37", "36")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    if title:
        run_title = p.add_run(f"{title}\n")
        run_title.bold = True
        run_title.font.name = 'Golos Text'
        run_title.font.size = Pt(11)
        run_title.font.color.rgb = RGBColor(5, 20, 42) # Primary Blue
        
    run_text = p.add_run(text)
    run_text.font.name = 'Golos Text'
    run_text.font.size = Pt(10.5)
    run_text.font.color.rgb = RGBColor(31, 41, 55) # Text Dark
    
    doc.add_paragraph() # Espaciado posterior

def build_brand_manual():
    doc = Document()

    # Configuración de márgenes estándar (1 pulgada = 2.54 cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Definición de Colores Corporativos
    COLOR_PRIMARY = RGBColor(5, 20, 42)     # #05142a - Azul Marino Profundo
    COLOR_ACCENT = RGBColor(212, 175, 55)   # #d4af37 - Dorado Metálico
    COLOR_DARK = RGBColor(3, 10, 21)        # #030a15 - Fondo Oscuro Noche
    COLOR_TEXT = RGBColor(31, 41, 55)       # #1F2937 - Gris Oscuro Texto
    COLOR_MUTED = RGBColor(100, 110, 120)   # Gris Secundario

    FONT_FAMILY = 'Golos Text'

    # --- ESTILOS DE PARRAFO Y ENCABEZADOS ---
    styles = doc.styles

    # Normal Style
    style_normal = styles['Normal']
    style_normal.font.name = FONT_FAMILY
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = COLOR_TEXT
    style_normal.paragraph_format.line_spacing = 1.2
    style_normal.paragraph_format.space_after = Pt(6)

    # Header 1
    h1 = styles['Heading 1']
    h1.font.name = FONT_FAMILY
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_PRIMARY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)

    # Header 2
    h2 = styles['Heading 2']
    h2.font.name = FONT_FAMILY
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_ACCENT
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    # Header 3
    h3 = styles['Heading 3']
    h3.font.name = FONT_FAMILY
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = COLOR_PRIMARY
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)

    # =========================================================================
    # PORTADA / COVER PAGE
    # =========================================================================
    p_cover_space = doc.add_paragraph()
    p_cover_space.paragraph_format.space_before = Pt(40)

    p_brand = doc.add_paragraph()
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_brand = p_brand.add_run("LISBOA LEGAL GROUP")
    r_brand.font.name = FONT_FAMILY
    r_brand.font.size = Pt(28)
    r_brand.font.bold = True
    r_brand.font.color.rgb = COLOR_PRIMARY

    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tag = p_tag.add_run("Liderazgo Legal. Impacto Global.")
    r_tag.font.name = FONT_FAMILY
    r_tag.font.size = Pt(14)
    r_tag.font.italic = True
    r_tag.font.color.rgb = COLOR_ACCENT
    p_tag.paragraph_format.space_after = Pt(40)

    # Línea Divisoria de Portada
    t_div = doc.add_table(rows=1, cols=1)
    t_div.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_div = t_div.cell(0, 0)
    cell_div.width = Inches(6.5)
    set_cell_background(cell_div, "D4AF37")
    cell_div.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell_div.paragraphs[0].paragraph_format.space_after = Pt(2)

    doc.add_paragraph().paragraph_format.space_after = Pt(40)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("MANUAL DE IDENTIDAD Y ESTRATEGIA DE MARCA")
    r_title.font.name = FONT_FAMILY
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Guía Oficial de Normas Visuales, Filosofía Corporativa y Lineamientos de Comunicación")
    r_sub.font.name = FONT_FAMILY
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = COLOR_MUTED
    p_sub.paragraph_format.space_after = Pt(120)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run("Guayaquil, Ecuador — Versión 1.0 (2026)\nDocumento Institucional Reservado")
    r_meta.font.name = FONT_FAMILY
    r_meta.font.size = Pt(9.5)
    r_meta.font.color.rgb = COLOR_MUTED

    doc.add_page_break()

    # =========================================================================
    # TABLA DE CONTENIDOS INTRODUCTORIA
    # =========================================================================
    doc.add_heading("Índice del Manual de Marca", level=1)
    
    table_toc = doc.add_table(rows=8, cols=2)
    table_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_toc.autofit = False
    
    toc_data = [
        ("Sección 1: Filosofía y Esencia Corporativa", "Misión, Visión, Propuesta de Valor, Pilares y Tono de Voz"),
        ("Sección 2: Arquitectura de Identidad Visual", "Isologo, Símbolo, Paleta de Colores, Tipografía y Texturas"),
        ("Sección 3: Portafolio de Áreas de Práctica", "Definición de Servicios y Cobertura Jurídica Especializada"),
        ("Sección 4: Herramientas Legales e Innovación Tech", "Calculadoras Digitales Abiertas e Inteligencia Artificial"),
        ("Sección 5: Perfiles del Equipo de Socios", "Trayectoria, Estudios y Especialidades del Cuadro Directivo"),
        ("Sección 6: Cultura, Ética y Responsabilidad Social", "Estándares Internacionales, Cero Tolerancia y Pro Bono"),
        ("Sección 7: Aplicaciones de Marca y Canales de Comunicación", "Web, Papelería, Redes Sociales y Directorio de Contacto"),
        ("Sección 8: Gobernanza y Uso Correcto de Marca", "Matriz de Do's & Don'ts y Políticas de Control Institucional")
    ]
    
    for idx, (col1, col2) in enumerate(toc_data):
        row = table_toc.rows[idx]
        cell1, cell2 = row.cells[0], row.cells[1]
        cell1.width = Inches(2.8)
        cell2.width = Inches(3.7)
        
        bg_color = "F4F6F9" if idx % 2 == 0 else "FFFFFF"
        set_cell_background(cell1, bg_color)
        set_cell_background(cell2, bg_color)
        set_cell_margins(cell1, 80, 80, 100, 100)
        set_cell_margins(cell2, 80, 80, 100, 100)
        
        p1 = cell1.paragraphs[0]
        r1 = p1.add_run(col1)
        r1.font.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = COLOR_PRIMARY
        
        p2 = cell2.paragraphs[0]
        r2 = p2.add_run(col2)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = COLOR_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # =========================================================================
    # SECCIÓN 1: FILOSOFÍA Y ESENCIA CORPORATIVA
    # =========================================================================
    doc.add_heading("Sección 1: Filosofía y Esencia Corporativa", level=1)

    p = doc.add_paragraph()
    p.add_run("Lisboa Legal Group ").bold = True
    p.add_run("es una firma jurídica de alto prestigio radicada en Guayaquil, Ecuador. Se distingue en el mercado legal por combinar rigor técnico contundente, visión preventiva de riesgos corporativos y el uso pionero de tecnologías avanzadas e inteligencia artificial aplicada al patrocinio legal.")

    doc.add_heading("1.1 Misión Corporativa", level=2)
    make_callout(doc, 
        "\"Brindar soluciones jurídicas integrales, preventivas y estratégicas que protejan con absoluta excelencia el patrimonio, la reputación y los derechos fundamentales de nuestros clientes, combinando la experiencia técnica de nuestros socios con innovación tecnológica de vanguardia.\"",
        "MISIÓN DE LISBOA LEGAL GROUP"
    )

    doc.add_heading("1.2 Visión Institucional", level=2)
    make_callout(doc,
        "\"Consolidarnos como el estudio jurídico líder y de mayor referencia en Guayaquil y Ecuador en litigación de alta complejidad, compliance preventivo y estructuración corporativa, siendo reconocidos por nuestro impacto global, estándares éticos inflexibles y la transformación digital de la práctica legal.\"",
        "VISIÓN DE LISBOA LEGAL GROUP"
    )

    doc.add_heading("1.3 Propuesta de Valor Diferenciadora", level=2)
    
    p = doc.add_paragraph()
    p.add_run("La propuesta de valor de Lisboa Legal Group se estructura sobre tres ejes fundamentales:")

    table_pv = doc.add_table(rows=4, cols=2)
    table_pv.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_pv.autofit = False
    
    headers_pv = ["Eje Estratégico", "Descripción y Valor para el Cliente"]
    hdr_cells = table_pv.rows[0].cells
    for i, title in enumerate(headers_pv):
        cell = hdr_cells[i]
        set_cell_background(cell, "05142A")
        set_cell_margins(cell, 100, 100, 120, 120)
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)
    hdr_cells[0].width = Inches(2.2)
    hdr_cells[1].width = Inches(4.3)

    pv_data = [
        ("Excelencia Jurídica", "Mantenemos los más altos estándares éticos y profesionales, ofreciendo soluciones técnicas de altísima complejidad con precisión y resultados contundentes en cada caso, asesoría o litigio."),
        ("Enfoque Estratégico", "Diseñamos estrategias legales a la medida que trascienden el ámbito puramente procesal. Entendemos el negocio y los objetivos de nuestros clientes para mitigar riesgos, maximizar recursos y asegurar victorias clave."),
        ("Tecnología Aplicada", "Implementamos herramientas digitales propias de cálculo legal, IA avanzada y estrictos protocolos de protección de datos para agilizar el análisis jurisprudencial y optimizar la defensa jurídica.")
    ]

    for idx, (col1, col2) in enumerate(pv_data):
        row = table_pv.rows[idx + 1]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width = Inches(2.2)
        c2.width = Inches(4.3)
        bg = "F4F6F9" if idx % 2 == 0 else "FFFFFF"
        set_cell_background(c1, bg)
        set_cell_background(c2, bg)
        set_cell_margins(c1, 80, 80, 100, 100)
        set_cell_margins(c2, 80, 80, 100, 100)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(col1)
        r1.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = COLOR_PRIMARY
        
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(col2)
        r2.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    doc.add_heading("1.4 Personalidad y Tono de Voz", level=2)
    p = doc.add_paragraph()
    p.add_run("El tono de voz de Lisboa Legal Group refleja sobriedad, firmeza, autoridad legal y modernidad. Toda pieza publicitaria, documento institucional o comunicación a clientes debe alinearse con estos principios:")

    bullet_points = [
        ("Riguroso y Sobrio: ", "Vocabulario preciso, sin tecnicismos innecesariamente confusos pero manteniendo la elegancia técnica del derecho."),
        ("Estratégico y Resolutivo: ", "Centrado siempre en la solución, la anticipación al conflicto y el blindaje patrimonial."),
        ("Accesible y Humano: ", "Especialmente en materias sensibles como el Derecho de Familia y Sucesiones, demostrando empatía y comprensión profesional."),
        ("Innovador y Sofisticado: ", "Uso de un lenguaje contemporáneo que proyecte liderazgo tecnológico y visión global.")
    ]

    for title, desc in bullet_points:
        bp = doc.add_paragraph(style='List Bullet')
        r_t = bp.add_run(title)
        r_t.bold = True
        r_t.font.color.rgb = COLOR_PRIMARY
        r_d = bp.add_run(desc)

    doc.add_page_break()

    # =========================================================================
    # SECCIÓN 2: ARQUITECTURA DE IDENTIDAD VISUAL
    # =========================================================================
    doc.add_heading("Sección 2: Arquitectura de Identidad Visual", level=1)

    doc.add_heading("2.1 Logotipo Principal (Isologo)", level=2)
    p = doc.add_paragraph()
    p.add_run("El identificador visual principal de la marca es ").font.color.rgb = COLOR_TEXT
    p.add_run("logo-full.png").bold = True
    p.add_run(". Consta del imagotipo estilizado acompañado de la tipografía corporativa institucional. Debe conservarse siempre la proporción de aspecto original sin distorsiones horizontales o verticales.")

    doc.add_heading("2.2 Paleta Cromática Oficial", level=2)
    p = doc.add_paragraph()
    p.add_run("La paleta de colores ha sido elegida minuciosamente para transmitir solidez institucional, sofisticación, confianza y prestigio.")

    table_color = doc.add_table(rows=6, cols=5)
    table_color.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_color.autofit = False

    c_headers = ["Nombre del Color", "Muestra Visual", "HEX", "RGB", "Función en la Marca"]
    for i, h in enumerate(c_headers):
        c = table_color.rows[0].cells[i]
        set_cell_background(c, "05142A")
        set_cell_margins(c, 100, 100, 100, 100)
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)

    color_data = [
        ("Azul Marino Profundo", "05142A", "#05142A", "5, 20, 42", "Color Primario Institucional / Encabezados"),
        ("Fondo Noche Oscuro", "030A15", "#030A15", "3, 10, 21", "Fondos de alto contraste y Hero Headers"),
        ("Dorado Metálico", "D4AF37", "#D4AF37", "212, 175, 55", "Color de Acento, Elegancia y Botones CTA"),
        ("Dorado Hover / Sec.", "B5952F", "#B5952F", "181, 149, 47", "Estados activos e interacciones de usuario"),
        ("Gris Oscuro Texto", "1F2937", "#1F2937", "31, 41, 55", "Cuerpo de texto y párrafos institucionales")
    ]

    widths = [Inches(1.5), Inches(1.0), Inches(0.9), Inches(1.0), Inches(2.1)]

    for idx, (name, hex_code, hex_str, rgb_str, role) in enumerate(color_data):
        row = table_color.rows[idx + 1]
        for c_idx in range(5):
            cell = row.cells[c_idx]
            cell.width = widths[c_idx]
            set_cell_margins(cell, 80, 80, 80, 80)
            
            p = cell.paragraphs[0]
            if c_idx == 0:
                r = p.add_run(name)
                r.bold = True
                r.font.size = Pt(9)
            elif c_idx == 1:
                set_cell_background(cell, hex_code)
            elif c_idx == 2:
                r = p.add_run(hex_str)
                r.font.size = Pt(9)
            elif c_idx == 3:
                r = p.add_run(rgb_str)
                r.font.size = Pt(9)
            elif c_idx == 4:
                r = p.add_run(role)
                r.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    doc.add_heading("2.3 Sistema Tipográfico", level=2)
    p = doc.add_paragraph()
    p.add_run("La tipografía oficial de Lisboa Legal Group es ").font.color.rgb = COLOR_TEXT
    p.add_run("Golos Text").bold = True
    p.add_run(", una familia tipográfica de grotesca moderna con una geometría limpia, leibilidad óptima y carácter corporativo contemporáneo.")

    table_typo = doc.add_table(rows=5, cols=3)
    table_typo.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_typo.autofit = False

    t_headers = ["Variante / Peso", "Tamaño / Uso", "Ejemplo de Aplicación"]
    for i, h in enumerate(t_headers):
        c = table_typo.rows[0].cells[i]
        set_cell_background(c, "05142A")
        set_cell_margins(c, 100, 100, 100, 100)
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9.5)

    typo_data = [
        ("Golos Text Bold (700)", "2.8rem / 28pt", "Titulares Principales H1 y Hero Headers"),
        ("Golos Text Semi-Bold (600)", "1.4rem - 1.8rem / 14-18pt", "Subtítulos H2, Nombre de Socios y Botones CTA"),
        ("Golos Text Medium (500)", "1.1rem / 11-12pt", "Destacados, Tarjetas de Servicios e Items de Menú"),
        ("Golos Text Regular (400)", "0.95rem - 1.05rem / 10-10.5pt", "Cuerpo de texto, descripciones y notas legales")
    ]

    for idx, (w, s, e) in enumerate(typo_data):
        row = table_typo.rows[idx + 1]
        for c_i, val in enumerate([w, s, e]):
            cell = row.cells[c_i]
            set_cell_margins(cell, 80, 80, 80, 80)
            bg = "F4F6F9" if idx % 2 == 0 else "FFFFFF"
            set_cell_background(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            if c_i == 0:
                r.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    doc.add_heading("2.4 Texturas y Elementos Gráficos de Marca", level=2)
    bp_graphics = [
        ("Textura Mármol Blanco (white-marble.png): ", "Utilizada como fondo sutil en secciones institucionales (.bg-marble) para transmitir lujo, nobleza, pulcritud y la solidez del derecho clásico."),
        ("Línea de Acento Dorada: ", "Líneas de 70px de ancho y 3px de alto en color Dorado Metálico (#d4af37) ubicadas debajo de los títulos de sección para jerarquizar la lectura."),
        ("Glassmorphic Header & Backdrop Blur: ", "Uso de transparencias con desenfoque de fondo en cabeceras de navegación (rgba(5, 20, 42, 0.95) con blur de 10px), proyectando modernidad y tecnología.")
    ]

    for title, desc in bp_graphics:
        bp = doc.add_paragraph(style='List Bullet')
        r_t = bp.add_run(title)
        r_t.bold = True
        r_t.font.color.rgb = COLOR_PRIMARY
        r_d = bp.add_run(desc)

    doc.add_page_break()

    # =========================================================================
    # SECCIÓN 3: PORTAFOLIO DE ÁREAS DE PRÁCTICA
    # =========================================================================
    doc.add_heading("Sección 3: Portafolio de Áreas de Práctica", level=1)

    p = doc.add_paragraph()
    p.add_run("Lisboa Legal Group ofrece asesoría legal especializada en 10 áreas fundamentales del derecho ecuatoriano e internacional:")

    services = [
        ("1. Derecho Penal", "Expertos en litigación penal estratégica con capacidad operativa a nivel nacional e internacional. Ofrecemos soluciones jurídicas de alta complejidad para los sectores público y privado, penal militar y blindaje ante delitos de cuello blanco."),
        ("2. Derecho Civil", "Especialistas en la protección del patrimonio, obligaciones contractuales y demandas por responsabilidad civil. Gestionamos controversias comerciales y civiles complejas con máxima rigurosidad."),
        ("3. Derecho de Familia", "Asesoramiento integral e interdisciplinario en procedimientos de divorcio, fijación e incremento de pensiones alimenticias, tenencia compartida y acuerdos prematrimoniales, velando por los derechos de la familia."),
        ("4. Derecho Constitucional", "Patrocinio de alto nivel en acciones de protección, hábeas corpus, hábeas data y acciones extraordinarias de protección frente a actuaciones u omisiones ilegítimas de autoridades públicas o privadas."),
        ("5. Derecho Administrativo", "Representación experta ante la Administración Pública. Asesoría en contratación pública (SERCOP), impugnación de actos administrativos, recursos en sede administrativa y defensas en procedimientos sancionadores y coactivos."),
        ("6. Corporate Compliance", "Diseño, auditoría e implementación de programas de cumplimiento normativo, mapas de riesgo corporativo y sistemas anticorrupción/prevención de lavado de activos para blindar el prestigio de corporaciones nacionales e internacionales."),
        ("7. Derecho de Sucesiones", "Planificación patrimonial y hereditaria a medida. Asistencia legal en testamentos, particiones de herencia extrajudiciales o judiciales, constitución de fideicomisos y solución de conflictos familiares complejos."),
        ("8. Derecho Laboral", "Estrategias laborales corporativas en contratación, auditorías de Nómina e IESS, prevención de riesgos patronales, manejo de desvinculaciones estratégicas y representación en litigios laborales individuales o colectivos."),
        ("9. Derecho Societario", "Estructuración corporativa integral: constitución de compañías, reformas estatutarias, fusiones, escisiones, disoluciones y gobierno corporativo para garantizar la solidez legal del negocio."),
        ("10. Consultoría Tributaria & Derecho Procesal", "Optimización fiscal preventiva, defensa contencioso-tributaria y manejo de técnicas de litigación en audiencias procesales bajo el Código Orgánico General de Procesos (COGEP).")
    ]

    table_serv = doc.add_table(rows=len(services)+1, cols=2)
    table_serv.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_serv.autofit = False

    s_hdr = table_serv.rows[0].cells
    s_hdr[0].width = Inches(2.2)
    s_hdr[1].width = Inches(4.3)
    set_cell_background(s_hdr[0], "05142A")
    set_cell_background(s_hdr[1], "05142A")
    set_cell_margins(s_hdr[0], 100, 100, 100, 100)
    set_cell_margins(s_hdr[1], 100, 100, 100, 100)

    r0 = s_hdr[0].paragraphs[0].add_run("Área de Práctica")
    r0.bold = True
    r0.font.color.rgb = RGBColor(255, 255, 255)
    r0.font.size = Pt(10)

    r1 = s_hdr[1].paragraphs[0].add_run("Alcance del Servicio y Cobertura Legal")
    r1.bold = True
    r1.font.color.rgb = RGBColor(255, 255, 255)
    r1.font.size = Pt(10)

    for idx, (s_name, s_desc) in enumerate(services):
        row = table_serv.rows[idx + 1]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.2)
        c1.width = Inches(4.3)
        bg = "F4F6F9" if idx % 2 == 0 else "FFFFFF"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_margins(c0, 80, 80, 100, 100)
        set_cell_margins(c1, 80, 80, 100, 100)
        
        p0 = c0.paragraphs[0]
        r_n = p0.add_run(s_name)
        r_n.bold = True
        r_n.font.size = Pt(9.5)
        r_n.font.color.rgb = COLOR_PRIMARY
        
        p1 = c1.paragraphs[0]
        r_d = p1.add_run(s_desc)
        r_d.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # =========================================================================
    # SECCIÓN 4: HERRAMIENTAS DIGITALES E INNOVACIÓN TECH
    # =========================================================================
    doc.add_heading("Sección 4: Herramientas Legales e Innovación Tech", level=1)

    p = doc.add_paragraph()
    p.add_run("Como parte del compromiso con la modernización de la abogacía, Lisboa Legal Group ha desarrollado e integrado un ecosistema de herramientas digitales abiertas para la ciudadanía y empresas en Ecuador:")

    tools = [
        ("Calculadora Laboral Automatizada", "Permite calcular con precisión liquidaciones por despido intempestivo, desahucio, décimos y vacaciones no gozadas conforme al Código del Trabajo ecuatoriano."),
        ("Calculadora de Plazos Procesales", "Herramienta automatizada para la determinación exacta de días hábiles, términos y plazos procesales en materias civil, penal, laboral y administrativa bajo la normativa COGEP/COIP."),
        ("Calculador de Pensiones Jubilares", "Calcula estimaciones oficiales de jubilación patronal y referencial IESS aplicando las tablas oficiales y coeficientes del año 2026."),
        ("Calculadora de Pensión de Alimentos", "Determina los valores oficiales de manutención infantil según la Tabla Nacional de Pensiones Alimenticias 2026 del MIES/SUPA, incorporando prorrateos e incrementos por discapacidad.")
    ]

    for t_name, t_desc in tools:
        bp = doc.add_paragraph(style='List Bullet')
        r_t = bp.add_run(f"{t_name}: ")
        r_t.bold = True
        r_t.font.color.rgb = COLOR_PRIMARY
        r_d = bp.add_run(t_desc)

    make_callout(doc,
        "La integración de Inteligencia Artificial aplicada en el análisis previo de jurisprudencia y encriptación de datos sensibles de clientes posiciona a Lisboa Legal Group a la vanguardia de la ciberseguridad y eficiencia legal en la región.",
        "TECNOLOGÍA APLICADA Y SEGURIDAD"
    )

    doc.add_page_break()

    # =========================================================================
    # SECCIÓN 5: EQUIPO LÍDER Y PERFILES PROFESIONALES
    # =========================================================================
    doc.add_heading("Sección 5: Equipo Líder y Perfiles Profesionales", level=1)

    p = doc.add_paragraph()
    p.add_run("La firma se fundamenta en la solvencia técnica, formación académica de posgrado y reputación intachable de sus socios directores:")

    partners = [
        ("Daniel Espinoza", "Socio Fundador", "Especialista Penal Militar y Litigación Estratégica. Lidera las áreas de Derecho Penal y Defensa Técnica en causas de alta complejidad nacional e internacional."),
        ("Karol Carvajal", "Socia Directora", "Especialista en Derecho Corporativo y Civil. Encabezó la estructuración de la práctica societaria, contratos mercantiles y la gestión de patrimonio empresarial."),
        ("Gustavo San Andrés", "Socio Director", "Abogado de los Tribunales y Juzgados de la República, Magíster en Derecho Constitucional y Mediador Certificado. Cuenta con más de 10 años de trayectoria (desde 2014) atendiendo a Empresas Privadas, Instituciones Públicas y Personas Naturales.")
    ]

    for p_name, p_role, p_bio in partners:
        doc.add_heading(f"{p_name} — {p_role}", level=2)
        p = doc.add_paragraph()
        p.add_run(p_bio)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # =========================================================================
    # SECCIÓN 6: CULTURA CORPORATIVA Y ÉTICA
    # =========================================================================
    doc.add_heading("Sección 6: Cultura Corporativa, Ética y Pro Bono", level=1)

    bullets_culture = [
        ("Certificaciones y Estándares de Calidad: ", "Procesos auditados bajo métricas de cumplimiento normativo y excelencia operacional."),
        ("Política de Tolerancia Cero: ", "Estricto protocolo anticorrupción y desestimación inmediata de asuntos que impliquen conflictos de interés."),
        ("Trabajo Pro Bono y Responsabilidad Social: ", "Patrocinio jurídico gratuito en casos seleccionados de alta vulnerabilidad para promover el acceso igualitario a la justicia.")
    ]

    for title, desc in bullets_culture:
        bp = doc.add_paragraph(style='List Bullet')
        r_t = bp.add_run(title)
        r_t.bold = True
        r_t.font.color.rgb = COLOR_PRIMARY
        r_d = bp.add_run(desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # =========================================================================
    # SECCIÓN 7: APLICACIONES DE MARCA Y CANALES DE CONTACTO
    # =========================================================================
    doc.add_heading("Sección 7: Aplicaciones de Marca y Contacto Oficial", level=1)

    doc.add_heading("7.1 Canales Digitales y Presencia de Marca", level=2)
    p = doc.add_paragraph()
    p.add_run("Toda plataforma o perfil oficial de Lisboa Legal Group debe mantener coherencia visual estricta con los parámetros de este manual:")

    table_contact = doc.add_table(rows=6, cols=2)
    table_contact.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_contact.autofit = False

    tc_hdr = table_contact.rows[0].cells
    tc_hdr[0].width = Inches(2.2)
    tc_hdr[1].width = Inches(4.3)
    set_cell_background(tc_hdr[0], "05142A")
    set_cell_background(tc_hdr[1], "05142A")
    set_cell_margins(tc_hdr[0], 100, 100, 100, 100)
    set_cell_margins(tc_hdr[1], 100, 100, 100, 100)

    tc_hdr[0].paragraphs[0].add_run("Canal / Plataforma").bold = True
    tc_hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    tc_hdr[1].paragraphs[0].add_run("Detalle Oficial de Contacto").bold = True
    tc_hdr[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    c_info = [
        ("Sede Matriz Guayaquil", "Edificio Finansa, Vélez 220 entre Chile y Chimborazo, Guayaquil - Ecuador"),
        ("Teléfono Directo / WhatsApp", "(+593) 9 90967952 / 0990967952"),
        ("Correo Electrónico Oficial", "lisboalegalgroup@gmail.com"),
        ("Página Web Institucional", "https://lisboalegalgroup.github.io/legal/"),
        ("Redes Sociales Oficiales", "Facebook (@lisboalegalgroup), Instagram (@lisboalegalgroup), LinkedIn (lisboa-legal-group)")
    ]

    for idx, (canal, val) in enumerate(c_info):
        row = table_contact.rows[idx + 1]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.2)
        c1.width = Inches(4.3)
        bg = "F4F6F9" if idx % 2 == 0 else "FFFFFF"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_margins(c0, 80, 80, 100, 100)
        set_cell_margins(c1, 80, 80, 100, 100)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(canal)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = COLOR_PRIMARY
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # =========================================================================
    # SECCIÓN 8: MATRIZ DE GOBERNANZA Y USOS PROHIBIDOS (DON'TS)
    # =========================================================================
    doc.add_heading("Sección 8: Matriz de Gobernanza y Usos Prohibidos", level=1)

    table_donts = doc.add_table(rows=5, cols=2)
    table_donts.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_donts.autofit = False

    td_hdr = table_donts.rows[0].cells
    td_hdr[0].width = Inches(3.25)
    td_hdr[1].width = Inches(3.25)
    set_cell_background(td_hdr[0], "05142A")
    set_cell_background(td_hdr[1], "05142A")
    set_cell_margins(td_hdr[0], 100, 100, 100, 100)
    set_cell_margins(td_hdr[1], 100, 100, 100, 100)

    r_do = td_hdr[0].paragraphs[0].add_run("Usos Permitidos (DO'S)")
    r_do.bold = True
    r_do.font.color.rgb = RGBColor(255, 255, 255)
    
    r_dont = td_hdr[1].paragraphs[0].add_run("Usos Prohibidos (DON'TS)")
    r_dont.bold = True
    r_dont.font.color.rgb = RGBColor(255, 255, 255)

    matrix_rules = [
        ("Usar el logo sobre fondos en contraste alto (Azul Marino #05142a o Blanco #FFFFFF).", "No estirar, comprimir ni deformar las proporciones originales del logo."),
        ("Respetar la paleta de colores oficial HEX y los pesos tipográficos de Golos Text.", "No alterar los colores corporativos usando tonos dorados brillantes o azules eléctricos desalineados."),
        ("Mantener el tono de voz profesional, preventivo y tecnológicamente vanguardista.", "No usar lenguaje informal desproporcionado o impreciso en publicaciones legales."),
        ("Utilizar los elementos visuales aprobados como la textura de mármol y barras doradas.", "No saturar las piezas gráficas con sombras excesivas, bordes multicolores o imágenes de baja calidad.")
    ]

    for idx, (do_rule, dont_rule) in enumerate(matrix_rules):
        row = table_donts.rows[idx + 1]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(3.25)
        c1.width = Inches(3.25)
        bg = "F4F6F9" if idx % 2 == 0 else "FFFFFF"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_margins(c0, 80, 80, 100, 100)
        set_cell_margins(c1, 80, 80, 100, 100)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(f"✓  {do_rule}")
        r0.font.size = Pt(9)
        r0.font.color.rgb = RGBColor(16, 124, 65) # Verde éxito
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(f"✗  {dont_rule}")
        r1.font.size = Pt(9)
        r1.font.color.rgb = RGBColor(180, 40, 40) # Rojo alerta

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    p_final = doc.add_paragraph()
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_final.add_run("© 2026 Lisboa Legal Group. Todos los derechos reservados.")
    r_f.font.size = Pt(9)
    r_f.font.color.rgb = COLOR_MUTED

    # Guardar documento Word
    output_filename = "Manual_de_Marca_Lisboa_Legal_Group.docx"
    output_path = os.path.join(r"c:\Users\ja_Ca\Desktop\lisboa legal group", output_filename)
    doc.save(output_path)
    print(f"Documento creado exitosamente en: {output_path}")

if __name__ == "__main__":
    build_brand_manual()
