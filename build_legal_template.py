import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Establece el color de fondo de una celda en Hexadecimal."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """Establece márgenes internos para celdas de tabla."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_left_border(cell, color_hex="D4AF37", size="36"):
    """Agrega un borde izquierdo grueso a una celda para resaltar áreas clave."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size)
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color_hex)
    tcBorders.append(left)
    
    for b_name in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{b_name}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
        
    tcPr.append(tcBorders)

def add_header_and_footer(doc, logo_path):
    """Configura el encabezado institucional con logo y el pie de página con contacto."""
    section = doc.sections[0]
    
    # --- ENCABEZADO ---
    header = section.header
    p_head = header.paragraphs[0]
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Crear una tabla invisible de 1 fila y 2 columnas en el encabezado
    htable = header.add_table(rows=1, cols=2, width=Inches(6.5))
    htable.alignment = WD_TABLE_ALIGNMENT.CENTER
    htable.autofit = False
    
    c_logo, c_info = htable.rows[0].cells[0], htable.rows[0].cells[1]
    c_logo.width = Inches(2.2)
    c_info.width = Inches(4.3)
    
    # Quitar bordes a la tabla del encabezado
    for cell in [c_logo, c_info]:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for b_name in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{b_name}')
            b.set(qn('w:val'), 'none')
            tcBorders.append(b)
        tcPr.append(tcBorders)
    
    # Insertar Logo en celda izquierda
    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_logo = p_logo.add_run()
    if os.path.exists(logo_path):
        r_logo.add_picture(logo_path, width=Inches(1.8))
        
    # Datos del Estudio en celda derecha
    p_info = c_info.paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_firm = p_info.add_run("LISBOA LEGAL GROUP\n")
    r_firm.bold = True
    r_firm.font.name = 'Golos Text'
    r_firm.font.size = Pt(9.5)
    r_firm.font.color.rgb = RGBColor(5, 20, 42)
    
    r_sub = p_info.add_run("Edificio Finansa, Vélez 220 entre Chile y Chimborazo | Guayaquil, Ecuador\nTel: (+593) 9 90967952 | Correo: lisboalegalgroup@gmail.com")
    r_sub.font.name = 'Golos Text'
    r_sub.font.size = Pt(8)
    r_sub.font.color.rgb = RGBColor(100, 110, 120)

    # --- PIE DE PÁGINA ---
    footer = section.footer
    p_foot = footer.paragraphs[0]
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run("— LISBOA LEGAL GROUP • Liderazgo Legal. Impacto Global. —\nDocumento Judicial Oficial")
    r_foot.font.name = 'Golos Text'
    r_foot.font.size = Pt(8.5)
    r_foot.font.italic = True
    r_foot.font.color.rgb = RGBColor(100, 110, 120)

def create_legal_template():
    doc = Document()
    
    # Configuración de Márgenes Judiciales Estándar
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.different_first_page_header_footer = False

    logo_path = r"c:\Users\ja_Ca\Desktop\lisboa legal group\logo-full.png"
    add_header_and_footer(doc, logo_path)

    # Definición de Colores de Marca
    COLOR_PRIMARY = RGBColor(5, 20, 42)     # #05142a - Azul Marino Profundo
    COLOR_ACCENT = RGBColor(212, 175, 55)   # #d4af37 - Dorado Metálico
    COLOR_TEXT = RGBColor(31, 41, 55)       # #1F2937 - Gris Oscuro
    COLOR_MUTED = RGBColor(100, 110, 120)

    FONT_FAMILY = 'Golos Text'

    # Estilos Base
    styles = doc.styles
    style_normal = styles['Normal']
    style_normal.font.name = FONT_FAMILY
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_TEXT
    style_normal.paragraph_format.line_spacing = 1.25
    style_normal.paragraph_format.space_after = Pt(6)

    # --- TABLA DE SUMILLA / DATOS DEL PROCESO ---
    t_sumilla = doc.add_table(rows=5, cols=2)
    t_sumilla.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_sumilla.autofit = False

    fields = [
        ("SEÑOR/A JUEZ/A:", "[UNIDAD JUDICIAL / TRIBUNAL COMPETENTE DE GUAYAQUIL]"),
        ("CAUSA / PROCESO N.°:", "[EJ. 09332-2026-00123]"),
        ("ACTOR / SOLICITANTE:", "[NOMBRES Y APELLIDOS DEL ACTOR / CLIENTE]"),
        ("DEMANDADO:", "[NOMBRES Y APELLIDOS DE LA PARTE DEMANDADA]"),
        ("ASUNTO / MATERIA:", "[DERECHO PENAL / CIVIL / LABORAL / FAMILIA / CONSTITUCIONAL]")
    ]

    for idx, (label, val) in enumerate(fields):
        row = t_sumilla.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.2)
        c1.width = Inches(4.3)
        
        bg = "05142A" if idx == 0 else ("F4F6F9" if idx % 2 == 1 else "FFFFFF")
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_margins(c0, 60, 60, 100, 100)
        set_cell_margins(c1, 60, 60, 100, 100)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = RGBColor(255, 255, 255) if idx == 0 else COLOR_PRIMARY
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)
        if idx == 0:
            r1.bold = True
            r1.font.color.rgb = RGBColor(255, 255, 255)

    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(12)

    # --- TÍTULO DEL ESCRITO ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("[SUMILLA: PRESENTA ESCRITO / DEMANDA / CONTESTACIÓN / IMPUGNACIÓN]")
    r_title.bold = True
    r_title.font.name = FONT_FAMILY
    r_title.font.size = Pt(13)
    r_title.font.color.rgb = COLOR_PRIMARY
    p_title.paragraph_format.space_after = Pt(14)

    # --- CUERPO DEL ESCRITO JUDICIAL ---

    # I. COMPARECENCIA
    h1 = doc.add_heading("I. COMPARECENCIA Y CALIDAD EN LA QUE ACTÚA", level=2)
    h1.runs[0].font.color.rgb = COLOR_PRIMARY
    h1.runs[0].font.size = Pt(12)

    p_comp = doc.add_paragraph()
    p_comp.paragraph_format.line_spacing = 1.25
    r_c1 = p_comp.add_run("Yo, ")
    r_c2 = p_comp.add_run("[NOMBRES COMPLETOS DEL CLIENTE O COMPARECIENTE]")
    r_c2.bold = True
    r_c3 = p_comp.add_run(", con cédula de ciudadanía / RUC N.° ")
    r_c4 = p_comp.add_run("[09XXXXXXXX]")
    r_c4.bold = True
    r_c5 = p_comp.add_run(", de estado civil [COMPLETAR], domiciliado/a en la ciudad de Guayaquil; comparezco ante su Autoridad dentro del proceso de la referencia y, bajo el patrocinio técnico de ")
    r_c6 = p_comp.add_run("LISBOA LEGAL GROUP")
    r_c6.bold = True
    r_c6.font.color.rgb = COLOR_PRIMARY
    r_c7 = p_comp.add_run(", manifiesto lo siguiente:")

    # II. ANTECEDENTES Y HECHOS
    h2 = doc.add_heading("II. ANTECEDENTES Y FUNDAMENTOS DE HECHO", level=2)
    h2.runs[0].font.color.rgb = COLOR_PRIMARY
    h2.runs[0].font.size = Pt(12)

    p_h1 = doc.add_paragraph(style='List Bullet')
    p_h1.add_run("PRIMERO: ").bold = True
    p_h1.add_run("[Describir el antecedente fáctico o procesal previo en orden cronológico conciso].")

    p_h2 = doc.add_paragraph(style='List Bullet')
    p_h2.add_run("SEGUNDO: ").bold = True
    p_h2.add_run("[Detallar las circunstancias específicas que motivan la presente solicitud legal].")

    p_h3 = doc.add_paragraph(style='List Bullet')
    p_h3.add_run("TERCERO: ").bold = True
    p_h3.add_run("[Exponer el impacto o necesidad jurídica concreta que fundamenta la pretensión].")

    # III. FUNDAMENTOS DE DERECHO
    h3 = doc.add_heading("III. FUNDAMENTOS DE DERECHO Y JURISPRUDENCIA", level=2)
    h3.runs[0].font.color.rgb = COLOR_PRIMARY
    h3.runs[0].font.size = Pt(12)

    p_d1 = doc.add_paragraph()
    p_d1.add_run("Fundamento la presente petición en las siguientes disposiciones legales vigentes en la República del Ecuador:")
    
    p_d_const = doc.add_paragraph(style='List Bullet')
    p_d_const.add_run("Constitución de la República del Ecuador: ").bold = True
    p_d_const.add_run("Artículo 75 (Derecho a la tutela judicial efectiva) y Artículo 76 (Garantías básicas del debido proceso).")

    p_d_ley = doc.add_paragraph(style='List Bullet')
    p_d_ley.add_run("Código Aplicable (COGEP / COIP / Código del Trabajo): ").bold = True
    p_d_ley.add_run("Artículo [COMPLETAR NÚMERO DE ARTÍCULO Y DISPOSICIÓN LEGAL ESPECÍFICA].")

    # Resaltado para Cita de Jurisprudencia
    t_quote = doc.add_table(rows=1, cols=1)
    t_quote.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_quote.autofit = False
    c_q = t_quote.cell(0, 0)
    c_q.width = Inches(6.5)
    set_cell_background(c_q, "F4F6F9")
    set_cell_left_border(c_q, "D4AF37", "36")
    set_cell_margins(c_q, 100, 100, 140, 140)

    p_q = c_q.paragraphs[0]
    rq_t = p_q.add_run("Criterio Jurisprudencial / Doctrina Relevante:\n")
    rq_t.bold = True
    rq_t.font.size = Pt(10)
    rq_t.font.color.rgb = COLOR_PRIMARY
    rq_b = p_q.add_run("\"[Insertar cita textual de la sentencia de la Corte Constitucional o Corte Nacional de Justicia que respalda la posición jurídica].\"")
    rq_b.font.italic = True
    rq_b.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # IV. PETICIÓN CONCRETA
    h4 = doc.add_heading("IV. PRETENTSIÓN Y PETICIÓN CONCRETA", level=2)
    h4.runs[0].font.color.rgb = COLOR_PRIMARY
    h4.runs[0].font.size = Pt(12)

    p_pet = doc.add_paragraph()
    p_pet.add_run("Con los antecedentes de hecho expuestos y el sólido amparo de derecho invocado, solicito formalmente a su Autoridad:")
    
    p_pet1 = doc.add_paragraph(style='List Bullet')
    p_pet1.add_run("1. ").bold = True
    p_pet1.add_run("[Formular de manera precisa la petición principal que el Juez debe resolver].")

    p_pet2 = doc.add_paragraph(style='List Bullet')
    p_pet2.add_run("2. ").bold = True
    p_pet2.add_run("[Formular peticiones secundarias, providencias o medidas cautelares si aplicare].")

    # V. ANEXOS Y PRUEBAS
    h5 = doc.add_heading("V. DOCUMENTOS ACOMPAÑADOS Y ANEXOS", level=2)
    h5.runs[0].font.color.rgb = COLOR_PRIMARY
    h5.runs[0].font.size = Pt(12)

    t_anexos = doc.add_table(rows=4, cols=3)
    t_anexos.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_anexos.autofit = False

    a_hdrs = ["Anexo", "Descripción del Documento / Prueba", "Naturaleza"]
    for i, h in enumerate(a_hdrs):
        c = t_anexos.rows[0].cells[i]
        set_cell_background(c, "05142A")
        set_cell_margins(c, 80, 80, 100, 100)
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)

    t_anexos.rows[0].cells[0].width = Inches(1.2)
    t_anexos.rows[0].cells[1].width = Inches(4.1)
    t_anexos.rows[0].cells[2].width = Inches(1.2)

    a_data = [
        ("Anexo 1", "Copia legible de cédula y papeleta de votación del compareciente", "Identificación"),
        ("Anexo 2", "Nombramiento / Poder de Procuración Judicial a favor del Abogado Patrocinador", "Habilitante"),
        ("Anexo 3", "[Insertar documento probatorio adicional: contrato, factura, peritaje, etc.]", "Documental")
    ]

    for idx, (anx, dsc, nat) in enumerate(a_data):
        row = t_anexos.rows[idx + 1]
        c0, c1, c2 = row.cells[0], row.cells[1], row.cells[2]
        c0.width, c1.width, c2.width = Inches(1.2), Inches(4.1), Inches(1.2)
        bg = "F4F6F9" if idx % 2 == 0 else "FFFFFF"
        for c in [c0, c1, c2]:
            set_cell_background(c, bg)
            set_cell_margins(c, 60, 60, 80, 80)
        
        c0.paragraphs[0].add_run(anx).bold = True
        c0.paragraphs[0].runs[0].font.size = Pt(9)
        c1.paragraphs[0].add_run(dsc).font.size = Pt(9)
        c2.paragraphs[0].add_run(nat).font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # VI. NOTIFICACIONES
    h6 = doc.add_heading("VI. NOTIFICACIONES Y CASILLA JUDICIAL", level=2)
    h6.runs[0].font.color.rgb = COLOR_PRIMARY
    h6.runs[0].font.size = Pt(12)

    p_notif = doc.add_paragraph()
    p_notif.add_run("Notificaciones que me correspondan las recibiré en los medios institucionales habilitados por el equipo de patrocinio jurídico de ")
    p_notif.add_run("LISBOA LEGAL GROUP").bold = True
    p_notif.add_run(":")

    bp_n1 = doc.add_paragraph(style='List Bullet')
    bp_n1.add_run("Casilla Judicial Electrónica Satje: ").bold = True
    bp_n1.add_run("[09XXXXXXXX / MATRÍCULA FORO DE ABOGADOS]")

    bp_n2 = doc.add_paragraph(style='List Bullet')
    bp_n2.add_run("Correos Electrónicos Oficiales: ").bold = True
    bp_n2.add_run("lisboalegalgroup@gmail.com / [CORREO DEL ABOGADO SOCIAL]")

    bp_n3 = doc.add_paragraph(style='List Bullet')
    bp_n3.add_run("Teléfono Directo / WhatsApp: ").bold = True
    bp_n3.add_run("(+593) 9 90967952")

    doc.add_paragraph().paragraph_format.space_after = Pt(25)

    # --- FIRMAS Y SUSCRIPCIÓN ---
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Es Justicia.— Guayaquil, ")
    r_sub.font.italic = True
    r_sub2 = p_sub.add_run("[DÍA] de [MES] de 2026.")
    r_sub2.font.italic = True
    p_sub.paragraph_format.space_after = Pt(45)

    # Tabla de Firmas (2 Firmas: Cliente + Abogado Patrocinador)
    t_signatures = doc.add_table(rows=2, cols=2)
    t_signatures.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_signatures.autofit = False

    c_sig1, c_sig2 = t_signatures.rows[0].cells[0], t_signatures.rows[0].cells[1]
    c_sig1.width = Inches(3.25)
    c_sig2.width = Inches(3.25)

    # Quitar bordes excepto línea superior para firma
    for c in [c_sig1, c_sig2]:
        tcPr = c._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '12')
        top.set(qn('w:color'), '05142A')
        tcBorders.append(top)
        for b_name in ['left', 'bottom', 'right']:
            b = OxmlElement(f'w:{b_name}')
            b.set(qn('w:val'), 'none')
            tcBorders.append(b)
        tcPr.append(tcBorders)

    p_s1 = c_sig1.paragraphs[0]
    p_s1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_s1.add_run("f.) _____________________________\n").bold = True
    p_s1.add_run("[NOMBRE DEL COMPARECIENTE]\n").bold = True
    p_s1.runs[1].font.size = Pt(9.5)
    p_s1.add_run("C.C. N.° [09XXXXXXXX]\nCompareciente").font.size = Pt(8.5)

    p_s2 = c_sig2.paragraphs[0]
    p_s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_s2.add_run("f.) _____________________________\n").bold = True
    p_s2.add_run("Abg. [DANIEL ESPINOZA / KAROL CARVAJAL / GUSTAVO SAN ANDRÉS]\n").bold = True
    p_s2.runs[1].font.size = Pt(9.5)
    p_s2.runs[1].font.color.rgb = COLOR_PRIMARY
    p_s2.add_run("LISBOA LEGAL GROUP\nMatrícula F.A. N.° [09-20XX-XXX]").font.size = Pt(8.5)

    # Guardar plantilla docx
    output_filename = "Plantilla_Escrito_Judicial_Lisboa_Legal_Group.docx"
    output_path = os.path.join(r"c:\Users\ja_Ca\Desktop\lisboa legal group", output_filename)
    doc.save(output_path)
    print(f"Plantilla creada exitosamente en: {output_path}")

if __name__ == "__main__":
    create_legal_template()
